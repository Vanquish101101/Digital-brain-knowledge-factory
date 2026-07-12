from docx import Document
from fpdf import FPDF

from kf.config import Settings
from kf.extract import extract_text


def _dummy_settings(**overrides) -> Settings:
    base = dict(
        postgres_host="localhost", postgres_port=5432, postgres_user="u",
        postgres_password="p", postgres_db="d", qdrant_url="http://localhost:6333",
        minio_endpoint="localhost:9000", minio_access_key="a", minio_secret_key="s",
        data_root="./data", model_cache_dir="./data/model-cache", embedding_model="m",
        openrouter_api_key="k", llm_model="l", ocr_languages="eng",
        image_caption_threshold_chars=20, vision_model="v",
        video_frame_interval_seconds=15, whisper_model_size="small",
        max_video_frames=20, synthesis_notes_dir="./notes",
    )
    base.update(overrides)
    return Settings(**base)


def test_extracts_plain_markdown(tmp_path):
    f = tmp_path / "note.md"
    f.write_text("# Заголовок\n\nТело заметки.", encoding="utf-8")

    assert extract_text(f, _dummy_settings()) == "# Заголовок\n\nТело заметки."


def test_extracts_csv_as_text(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("a,b\n1,2", encoding="utf-8")

    assert extract_text(f, _dummy_settings()) == "a,b\n1,2"


def test_extracts_docx_paragraphs(tmp_path):
    f = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("Второй абзац.")
    doc.save(f)

    text = extract_text(f, _dummy_settings())

    assert "Первый абзац." in text
    assert "Второй абзац." in text


def test_extracts_pdf_text(tmp_path):
    f = tmp_path / "doc.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="Hello from test PDF")
    pdf.output(str(f))

    text = extract_text(f, _dummy_settings())

    assert "Hello from test PDF" in text


def test_dispatches_image_to_ocr(tmp_path, monkeypatch):
    f = tmp_path / "shot.png"
    f.write_bytes(b"fakepng")
    monkeypatch.setattr("kf.extract.extract_text_from_image", lambda path, languages: "OCR TEXT HERE")

    text = extract_text(f, _dummy_settings(image_caption_threshold_chars=5))

    assert text == "OCR TEXT HERE"


def test_falls_back_to_caption_when_ocr_text_is_short(tmp_path, monkeypatch):
    f = tmp_path / "photo.jpg"
    f.write_bytes(b"fakejpg")
    monkeypatch.setattr("kf.extract.extract_text_from_image", lambda path, languages: "")
    monkeypatch.setattr("kf.extract.caption_image", lambda settings, path: "Фото заката над морем")

    text = extract_text(f, _dummy_settings(image_caption_threshold_chars=20))

    assert "Фото заката над морем" in text


def test_dispatches_video_to_transcript_and_frames(tmp_path, monkeypatch):
    f = tmp_path / "clip.mp4"
    f.write_bytes(b"fakevideo")
    fake_audio = tmp_path / "audio.wav"
    fake_frames_dir = tmp_path / "frames"
    fake_frames_dir.mkdir()
    fake_frame = fake_frames_dir / "frame_0000.png"
    fake_frame.write_bytes(b"fakeframe")

    monkeypatch.setattr("kf.extract.extract_audio", lambda path: fake_audio)
    monkeypatch.setattr(
        "kf.extract.transcribe_audio", lambda path, model_size, cache_dir: "Привет мир"
    )
    monkeypatch.setattr(
        "kf.extract.sample_frames", lambda path, interval_seconds, max_frames: (fake_frames_dir, [fake_frame])
    )
    monkeypatch.setattr(
        "kf.extract.extract_text_from_image", lambda path, languages: "текст на кадре"
    )

    text = extract_text(f, _dummy_settings(image_caption_threshold_chars=5))

    assert "[Транскрипт]" in text
    assert "Привет мир" in text
    assert "[Кадр 00:00]" in text
    assert "текст на кадре" in text


def test_extract_video_cleans_up_temp_audio_and_frames(tmp_path, monkeypatch):
    f = tmp_path / "clip.mp4"
    f.write_bytes(b"fakevideo")
    fake_audio = tmp_path / "audio.wav"
    fake_audio.write_bytes(b"fakeaudio")
    fake_frames_dir = tmp_path / "frames"
    fake_frames_dir.mkdir()
    fake_frame = fake_frames_dir / "frame_0000.png"
    fake_frame.write_bytes(b"fakeframe")

    monkeypatch.setattr("kf.extract.extract_audio", lambda path: fake_audio)
    monkeypatch.setattr(
        "kf.extract.transcribe_audio", lambda path, model_size, cache_dir: "Привет мир"
    )
    monkeypatch.setattr(
        "kf.extract.sample_frames", lambda path, interval_seconds, max_frames: (fake_frames_dir, [fake_frame])
    )
    monkeypatch.setattr(
        "kf.extract.extract_text_from_image", lambda path, languages: "текст на кадре"
    )

    extract_text(f, _dummy_settings(image_caption_threshold_chars=5))

    assert not fake_audio.exists()
    assert not fake_frames_dir.exists()


def test_extract_video_cleans_up_temp_files_even_on_error(tmp_path, monkeypatch):
    f = tmp_path / "clip.mp4"
    f.write_bytes(b"fakevideo")
    fake_audio = tmp_path / "audio.wav"
    fake_audio.write_bytes(b"fakeaudio")
    fake_frames_dir = tmp_path / "frames"
    fake_frames_dir.mkdir()
    fake_frame = fake_frames_dir / "frame_0000.png"
    fake_frame.write_bytes(b"fakeframe")

    def _boom(path, languages):
        raise RuntimeError("OCR blew up")

    monkeypatch.setattr("kf.extract.extract_audio", lambda path: fake_audio)
    monkeypatch.setattr(
        "kf.extract.transcribe_audio", lambda path, model_size, cache_dir: "Привет мир"
    )
    monkeypatch.setattr(
        "kf.extract.sample_frames", lambda path, interval_seconds, max_frames: (fake_frames_dir, [fake_frame])
    )
    monkeypatch.setattr("kf.extract.extract_text_from_image", _boom)

    try:
        extract_text(f, _dummy_settings(image_caption_threshold_chars=5))
        raised = False
    except RuntimeError:
        raised = True

    assert raised
    assert not fake_audio.exists()
    assert not fake_frames_dir.exists()


def test_extract_video_passes_max_frames_cap_to_sample_frames(tmp_path, monkeypatch):
    f = tmp_path / "clip.mp4"
    f.write_bytes(b"fakevideo")
    fake_audio = tmp_path / "audio.wav"
    fake_frames_dir = tmp_path / "frames"
    fake_frames_dir.mkdir()

    received = {}

    def _fake_sample_frames(path, interval_seconds, max_frames):
        received["max_frames"] = max_frames
        return fake_frames_dir, []

    monkeypatch.setattr("kf.extract.extract_audio", lambda path: fake_audio)
    monkeypatch.setattr(
        "kf.extract.transcribe_audio", lambda path, model_size, cache_dir: "Привет мир"
    )
    monkeypatch.setattr("kf.extract.sample_frames", _fake_sample_frames)

    extract_text(f, _dummy_settings(max_video_frames=7))

    assert received["max_frames"] == 7
