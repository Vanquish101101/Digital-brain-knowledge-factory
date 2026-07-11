from docx import Document
from fpdf import FPDF

from kf.extract import extract_text


def test_extracts_plain_markdown(tmp_path):
    f = tmp_path / "note.md"
    f.write_text("# Заголовок\n\nТело заметки.", encoding="utf-8")

    assert extract_text(f) == "# Заголовок\n\nТело заметки."


def test_extracts_csv_as_text(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("a,b\n1,2", encoding="utf-8")

    assert extract_text(f) == "a,b\n1,2"


def test_extracts_docx_paragraphs(tmp_path):
    f = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("Второй абзац.")
    doc.save(f)

    text = extract_text(f)

    assert "Первый абзац." in text
    assert "Второй абзац." in text


def test_extracts_pdf_text(tmp_path):
    f = tmp_path / "doc.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="Hello from test PDF")
    pdf.output(str(f))

    text = extract_text(f)

    assert "Hello from test PDF" in text
